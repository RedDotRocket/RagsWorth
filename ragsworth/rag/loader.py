"""
Document loader implementations for various file types.
"""

import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Optional

from pypdf import PdfReader
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from ..pipeline import Document
from ..config.logging_config import get_logger

logger = get_logger("rag.loader")

class BaseLoader(ABC):
    """Base class for document loaders."""

    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        """Load documents from a file."""
        pass

class TextLoader(BaseLoader):
    """Loader for plain text files."""

    def load(self, file_path: str) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return [Document(
            id=str(Path(file_path).stem),
            content=content,
            metadata={
                "source": file_path,
                "type": "text"
            }
        )]

class MarkdownLoader(BaseLoader):
    """Loader for Markdown files."""

    def load(self, file_path: str) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return [Document(
            id=str(Path(file_path).stem),
            content=content,
            metadata={
                "source": file_path,
                "type": "markdown"
            }
        )]

class PDFLoader(BaseLoader):
    """Loader for PDF files."""

    def load(self, file_path: str) -> List[Document]:
        documents = []

        with open(file_path, 'rb') as f:
            pdf = PdfReader(f)

            for i, page in enumerate(pdf.pages):
                content = page.extract_text()
                if content.strip():
                    documents.append(Document(
                        id=f"{Path(file_path).stem}_page_{i+1}",
                        content=content,
                        metadata={
                            "source": file_path,
                            "type": "pdf",
                            "page": i + 1
                        }
                    ))

        return documents

class HTMLLoader(BaseLoader):
    """Loader for HTML files."""

    def load(self, file_path: str) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text content
            content = soup.get_text(separator='\n')

        return [Document(
            id=str(Path(file_path).stem),
            content=content,
            metadata={
                "source": file_path,
                "type": "html",
                "title": soup.title.string if soup.title else None
            }
        )]

class DocxLoader(BaseLoader):
    """Loader for Word documents."""

    def load(self, file_path: str) -> List[Document]:
        doc = DocxDocument(file_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        return [Document(
            id=str(Path(file_path).stem),
            content=content,
            metadata={
                "source": file_path,
                "type": "docx"
            }
        )]

class DocumentLoader:
    """Main document loader that handles different file types."""

    def __init__(self):
        self.loaders = {
            '.txt': TextLoader(),
            '.md': MarkdownLoader(),
            '.pdf': PDFLoader(),
            '.html': HTMLLoader(),
            '.htm': HTMLLoader(),
            '.docx': DocxLoader()
        }

    def load(self, file_path: str) -> List[Document]:
        """Load documents from a file based on its extension."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in self.loaders:
            raise ValueError(f"Unsupported file type: {ext}")

        return self.loaders[ext].load(file_path)

    def load_directory(self, directory: str, recursive: bool = True) -> List[Document]:
        """Load all supported documents from a directory."""
        documents = []

        for root, _, files in os.walk(directory):
            if not recursive and root != directory:
                continue

            for file in files:
                file_path = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()

                if ext in self.loaders:
                    try:
                        documents.extend(self.load(file_path))
                    except Exception as e:
                        logger.error(f"Error loading {file_path}: {e}", exc_info=True)

        return documents